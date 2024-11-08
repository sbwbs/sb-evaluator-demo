import streamlit as st
import pandas as pd
import json
from io import StringIO, BytesIO
import docx
from pathlib import Path
import tempfile
from collections import Counter
import re
from openai import OpenAI
from pydantic import BaseModel
from datetime import datetime

# Configure Streamlit page
st.set_page_config(page_title="Clinical Translation Evaluator", layout="wide")

# Initialize session state for API key
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = None

# Initialize OpenAI client
def get_openai_client(api_key):
    return OpenAI(api_key=api_key)

class ResponseEvaluation(BaseModel):
    status: str  # "PASS", "FAIL", "VERIFY"
    notes: str   # Additional notes for human QA if needed

def has_numerical_content(text):
    """Check if text contains any numbers or dates"""
    number_pattern = r'\d+\.?\d*'
    date_pattern = r'\b\d{1,4}[-/.]\d{1,2}[-/.]\d{1,4}\b'
    unit_pattern = r'\d+\s*(?:mg|kg|ml|g|mcg|µg|mmol|µmol|mmHg|°C|mm|cm|m|%)'
    
    return bool(re.search(number_pattern, str(text)) or 
               re.search(date_pattern, str(text)) or 
               re.search(unit_pattern, str(text)))

def evaluate_translation(source_text, target_text, client):
    """Evaluate translation using OpenAI API"""
    try:
        # Handle nan/NA cases first
        if pd.isna(source_text) or pd.isna(target_text) or str(target_text).lower() in ['na', 'nan']:
            return {"status": "PASS", "notes": "Target contains NA/nan or is empty - automatic PASS"}
        
        # Check if either segment contains numerical content
        if not has_numerical_content(source_text) and not has_numerical_content(target_text):
            return {"status": "PASS", "notes": "No numerical content to evaluate - automatic PASS"}

        messages = [
            {
                "role": "system",
                "content": f"""
                You are a clinical translation quality evaluator specialized in detecting ONLY numerical and temporal inconsistencies.
                
                EVALUATION SCOPE:
                ✓ Numbers (dosages, measurements, patient counts, percentages)
                ✓ Dates (study periods, follow-up times, patient visits)
                ✓ Clinical measurements/units
                ✓ Laboratory values
                
                IMPORTANT RULES:
                1. If there are NO numbers, dates, or measurements in either segment, return PASS
                2. Only evaluate numerical/date/measurement differences
                3. Ignore all other translation aspects
                4. If target contains NA/nan/na, return PASS
                
                Source Segment: {source_text}
                Target Segment: {target_text}
                
                Respond ONLY with a JSON in this exact format:
                {{
                    "status": "PASS|FAIL|VERIFY",
                    "notes": "Brief explanation focusing ONLY on numerical differences if found"
                }}
                """
            }
        ]

        completion = client.beta.chat.completions.parse(
            model="gpt-4o-2024-08-06",
            messages=messages,
            response_format=ResponseEvaluation,
            timeout=30  # OpenAI client timeout
        )

        response_content = json.loads(completion.choices[0].message.content)
        return {"status": response_content["status"], "notes": response_content["notes"]}

    except Exception as e:
        st.error(f"Error during evaluation: {str(e)}")
        return {"status": "VERIFY", "notes": f"Error encountered: {str(e)}. Requires human QA."}

# Also update the docx extraction function to handle both column naming conventions
def extract_table_from_docx(file_content):
    """Extract table content from DOCX file and convert to DataFrame"""
    doc = docx.Document(BytesIO(file_content))
    
    # Initialize lists to store table data
    data = []
    
    # Process the first table found in the document
    if doc.tables:
        table = doc.tables[0]  # Get the first table
        
        # Get headers from the first row
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())
        
        # Get data from remaining rows
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            data.append(row_data)
        
        # Create DataFrame
        df = pd.DataFrame(data, columns=headers)
        
        # Define possible column names
        source_columns = ['Source text', 'Source segment']
        target_columns = ['Target text', 'Target segment']
        
        # Check if at least one source and one target column exists
        found_source = [col for col in source_columns if col in df.columns]
        found_target = [col for col in target_columns if col in df.columns]
        
        if not found_source or not found_target:
            st.error(f"""Required columns not found in the document table. Please ensure it contains either:
            - 'Source text' and 'Target text' columns, OR
            - 'Source segment' and 'Target segment' columns""")
            return None
            
        # Use the first found column names
        source_col = found_source[0]
        target_col = found_target[0]
        
        # Rename columns to standardized names
        df = df.rename(columns={
            source_col: 'Source text',
            target_col: 'Target text'
        })
        
        return df
    else:
        st.error("No table found in the document. Please ensure the document contains a table with the required columns.")
        return None

def process_document(file_obj, file_type):
    """Process uploaded document based on file type"""
    try:
        if file_type == "csv":
            df = pd.read_csv(file_obj)
        elif file_type == "xlsx":
            df = pd.read_excel(file_obj)
        elif file_type == "docx":
            df = extract_table_from_docx(file_obj.getvalue())
            if df is None:
                return None
        
        # Define possible column names
        source_columns = ['Source text', 'Source segment']
        target_columns = ['Target text', 'Target segment']
        
        # Check if at least one source and one target column exists
        found_source = [col for col in source_columns if col in df.columns]
        found_target = [col for col in target_columns if col in df.columns]
        
        if not found_source or not found_target:
            st.error(f"""Required columns not found. Please ensure your file contains either:
            - 'Source text' and 'Target text' columns, OR
            - 'Source segment' and 'Target segment' columns""")
            return None
        
        # Use the first found column names
        source_col = found_source[0]
        target_col = found_target[0]
        
        # Rename columns to standardized names if necessary
        df = df.rename(columns={
            source_col: 'Source text',
            target_col: 'Target text'
        })
        
        return df
    
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def main():
    st.title("Clinical Translation Evaluation Tool")
    st.markdown("""
    This tool evaluates clinical translations for numerical and temporal consistency.
    Enter your OpenAI API key and upload your file to begin.
    """)
    
    # Initialize all session states
    if 'results' not in st.session_state:
        st.session_state.results = None
    if 'processed' not in st.session_state:
        st.session_state.processed = False
    if 'openai_api_key' not in st.session_state:
        st.session_state.openai_api_key = None

    # Add a clear button at the top
    if st.button("Clear and Start Over"):
        st.session_state.clear()
        st.rerun()
    
    # API Key input
    api_key = st.text_input("Enter your OpenAI API key:", type="password")
    if api_key:
        st.session_state.openai_api_key = api_key
    
    # Only show file upload and processing if not yet processed
    if not st.session_state.processed:
        # File upload section
        uploaded_file = st.file_uploader(
            "Upload translation file (must contain 'Source text'/'Target text' or 'Source segment'/'Target segment' columns)",
            type=["csv", "xlsx", "docx"],
            help="Supported formats: CSV, Excel, Word (with table)"
        )
        
        # Process control options
        test_mode = st.checkbox("Test mode (process first 50 rows only)", value=True)
        
        if uploaded_file and st.session_state.openai_api_key:
            client = get_openai_client(st.session_state.openai_api_key)
            file_type = uploaded_file.name.split('.')[-1].lower()
            
            with st.spinner('Processing translations...'):
                # Load and process the file
                df = process_document(uploaded_file, file_type)
                
                if df is None:
                    st.stop()  # Stop execution if file processing failed
                
                if test_mode:
                    df = df.head(50)
                
                # Display sample of data
                st.subheader("Sample of uploaded data")
                st.dataframe(df.head(), use_container_width=True)
                
                # Process translations
                results = []
                fail_ids = []
                
                progress_text = "Evaluating translations..."
                progress_bar = st.progress(0, text=progress_text)
                
                for index, row in df.iterrows():
                    source_text = row['Source text']
                    target_text = row['Target text']
                    
                    evaluation = evaluate_translation(source_text, target_text, client)
                    results.append(evaluation['status'])
                    
                    if evaluation['status'] == 'FAIL':
                        fail_ids.append(index)
                    
                    df.at[index, 'Evaluation'] = json.dumps(evaluation)
                    progress_bar.progress((index + 1) / len(df), 
                                       text=f"{progress_text} ({index + 1}/{len(df)})")
                
                # Calculate and store results
                result_counts = Counter(results)
                # Prepare download data once and store in session state
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                buffer = BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
                
                st.session_state.results = {
                    'df': df,
                    'counts': result_counts,
                    'fail_ids': fail_ids,
                    'total': len(results),
                    'excel_data': buffer.getvalue(),
                    'csv_data': df.to_csv(index=False),
                    'timestamp': timestamp
                }
                st.session_state.processed = True
                st.rerun()

    # Display results if available
    if st.session_state.results:
        st.markdown("## Evaluation Results")
        
        # Display metrics
        cols = st.columns(4)
        metrics = {
            "PASS": "green",
            "FAIL": "red",
            "VERIFY": "orange",
            "TOTAL": "blue"
        }
        
        for col, (metric, color) in zip(cols, metrics.items()):
            value = (st.session_state.results['counts'][metric] 
                    if metric != "TOTAL" 
                    else st.session_state.results['total'])
            col.metric(
                metric,
                value,
                delta=f"{value/st.session_state.results['total']*100:.1f}%" if value > 0 else "0%"
            )
        
        # Get segments that need attention (FAIL or VERIFY)
        attention_needed_df = st.session_state.results['df'][
            st.session_state.results['df']['Evaluation'].apply(
                lambda x: json.loads(x)['status'] in ['FAIL', 'VERIFY']
            )
        ]
        
        if not attention_needed_df.empty:
            st.markdown("### Segments Needing Attention")
            
            # Add tabs for different views
            tabs = st.tabs(["All Issues", "Failed Segments", "Segments Needing Verification"])
            
            with tabs[0]:
                st.markdown("#### All Segments Needing Review")
                display_df = attention_needed_df.copy()
                display_df['Status'] = display_df['Evaluation'].apply(
                    lambda x: json.loads(x)['status']
                )
                display_df['Notes'] = display_df['Evaluation'].apply(
                    lambda x: json.loads(x)['notes']
                )
                st.dataframe(
                    display_df[['Source text', 'Target text', 'Status', 'Notes']], 
                    use_container_width=True
                )
            
            with tabs[1]:
                st.markdown("#### Failed Segments")
                failed_df = attention_needed_df[
                    attention_needed_df['Evaluation'].apply(
                        lambda x: json.loads(x)['status'] == 'FAIL'
                    )
                ]
                if not failed_df.empty:
                    display_df = failed_df.copy()
                    display_df['Notes'] = display_df['Evaluation'].apply(
                        lambda x: json.loads(x)['notes']
                    )
                    st.dataframe(
                        display_df[['Source text', 'Target text', 'Notes']], 
                        use_container_width=True
                    )
                else:
                    st.info("No failed segments found.")
            
            with tabs[2]:
                st.markdown("#### Segments Needing Verification")
                verify_df = attention_needed_df[
                    attention_needed_df['Evaluation'].apply(
                        lambda x: json.loads(x)['status'] == 'VERIFY'
                    )
                ]
                if not verify_df.empty:
                    display_df = verify_df.copy()
                    display_df['Notes'] = display_df['Evaluation'].apply(
                        lambda x: json.loads(x)['notes']
                    )
                    st.dataframe(
                        display_df[['Source text', 'Target text', 'Notes']], 
                        use_container_width=True
                    )
                else:
                    st.info("No segments needing verification found.")
        else:
            st.success("No segments need attention - all translations PASSED! 🎉")
        
        # Download buttons
        st.markdown("### Download Results")
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                label="Download CSV",
                data=st.session_state.results['csv_data'],
                file_name=f"evaluated_translations_{st.session_state.results['timestamp']}.csv",
                mime="text/csv"
            )
        
        with col2:
            st.download_button(
                label="Download Excel",
                data=st.session_state.results['excel_data'],
                file_name=f"evaluated_translations_{st.session_state.results['timestamp']}.xlsx",
                mime="application/vnd.ms-excel"
            )

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        if st.button("Clear and Start Over"):
            st.session_state.clear()
            st.rerun()